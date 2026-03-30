#!/usr/bin/env python3
"""
Parse master_resume.docx into master_resume.json.

Usage:
    python scripts/parse_resume.py

Reads:  input/master_resume.docx
Writes: input/master_resume.json
"""

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
INPUT_DOCX = ROOT / "input" / "master_resume.docx"
OUTPUT_JSON = ROOT / "input" / "master_resume.json"

# ── Section heading detection ──────────────────────────────────────────────────
# Adjust these keywords to match your actual resume headings.
SECTION_KEYWORDS = [
    "summary", "objective", "profile",
    "experience", "work experience", "employment",
    "education",
    "skills", "technical skills", "core competencies",
    "projects",
    "certifications", "certificates", "licenses",
    "awards", "honors",
    "publications",
    "languages",
    "volunteer", "volunteering",
    "interests", "hobbies",
    "references",
]


def is_heading(paragraph) -> bool:
    """Return True if the paragraph looks like a section heading."""
    style_name = paragraph.style.name.lower()
    if "heading" in style_name:
        return True
    # Bold, short, all-caps lines are usually headings too
    text = paragraph.text.strip()
    if not text:
        return False
    if len(text) > 60:
        return False
    if text.lower() in SECTION_KEYWORDS:
        return True
    if text.upper() == text and len(text) > 2:
        return True
    # Check if every run is bold
    runs = [r for r in paragraph.runs if r.text.strip()]
    if runs and all(r.bold for r in runs):
        return True
    return False


def para_text(paragraph) -> str:
    return paragraph.text.strip()


def extract_contact(paragraphs: list) -> dict:
    """
    Heuristically pull name + contact info from the top of the document
    (before the first real section heading).
    """
    contact = {"name": "", "email": "", "phone": "", "linkedin": "", "location": "", "other": []}
    email_re = re.compile(r"[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}")
    phone_re = re.compile(r"[\+\(]?[\d][\d\s\-\(\)\.]{6,}[\d]")
    linkedin_re = re.compile(r"linkedin\.com/[\w/\-]+", re.I)

    for p in paragraphs:
        text = para_text(p)
        if not text:
            continue
        if not contact["name"]:
            # First non-empty paragraph before headings is the name
            contact["name"] = text
            continue
        if email_re.search(text):
            contact["email"] = email_re.search(text).group()
        if phone_re.search(text):
            contact["phone"] = phone_re.search(text).group().strip()
        if linkedin_re.search(text):
            contact["linkedin"] = linkedin_re.search(text).group()
        # Anything that looks like a city/state goes to location
        if re.search(r"\b[A-Z][a-z]+,\s*[A-Z]{2}\b", text):
            contact["location"] = text
    return contact


def table_to_text(table) -> list[str]:
    """Flatten a table into a list of non-empty cell strings."""
    rows = []
    for row in table.rows:
        for cell in row.cells:
            t = cell.text.strip()
            if t:
                rows.append(t)
    return rows


def parse_docx(path: Path) -> dict:
    doc = Document(path)

    # Collect all block-level content in document order (paragraphs + tables)
    blocks = []
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            # Find the matching Paragraph object
            for p in doc.paragraphs:
                if p._element is child:
                    blocks.append(("para", p))
                    break
        elif tag == "tbl":
            for t in doc.tables:
                if t._element is child:
                    blocks.append(("table", t))
                    break

    # ── Find where first section heading is ───────────────────────────────────
    first_heading_idx = None
    for i, (kind, obj) in enumerate(blocks):
        if kind == "para" and is_heading(obj):
            first_heading_idx = i
            break

    # Contact info lives before the first heading
    contact_blocks = blocks[:first_heading_idx] if first_heading_idx is not None else []
    contact_paras = [obj for kind, obj in contact_blocks if kind == "para"]
    contact = extract_contact(contact_paras)

    # ── Walk sections ──────────────────────────────────────────────────────────
    sections: dict[str, list[str]] = {}
    current_section: str | None = None

    content_blocks = blocks[first_heading_idx:] if first_heading_idx is not None else blocks

    for kind, obj in content_blocks:
        if kind == "para":
            text = para_text(obj)
            if not text:
                continue
            if is_heading(obj):
                current_section = text.lower()
                if current_section not in sections:
                    sections[current_section] = []
            else:
                if current_section is not None:
                    sections[current_section].append(text)
        elif kind == "table":
            if current_section is not None:
                sections[current_section].extend(table_to_text(obj))
            else:
                # Table before any heading — fold into contact
                contact["other"].extend(table_to_text(obj))

    return {"contact": contact, "sections": sections}


def main():
    if not INPUT_DOCX.exists():
        print(f"ERROR: {INPUT_DOCX} not found.", file=sys.stderr)
        print("Drop your master resume into input/master_resume.docx and re-run.", file=sys.stderr)
        sys.exit(1)

    print(f"Parsing {INPUT_DOCX} ...")
    data = parse_docx(INPUT_DOCX)

    OUTPUT_JSON.write_text(json.dumps(data, indent=2, ensure_ascii=False))
    print(f"Written → {OUTPUT_JSON}")

    # Quick summary
    print(f"  Name   : {data['contact'].get('name')}")
    print(f"  Email  : {data['contact'].get('email')}")
    print(f"  Sections found: {list(data['sections'].keys())}")


if __name__ == "__main__":
    main()
