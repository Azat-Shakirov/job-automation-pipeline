#!/usr/bin/env python3
"""
build_docs.py — build ATS-clean resume and cover letter .docx files.

Usage (CLI):
    python build_docs.py <resume_json> <company_name> "<cover letter text>"

Importable API:
    from build_docs import build_resume, build_cover_letter
"""

import json
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips

# ── Constants ──────────────────────────────────────────────────────────────────
# Sizes in actual points (XML w:sz stores half-points, so Pt(11) → w:sz="22")
FONT               = "Arial"
FONT_SIZE_NAME     = Pt(11)    # w:sz=22
FONT_SIZE_CONTACT  = Pt(8.5)   # w:sz=17
FONT_SIZE_HEADING  = Pt(9.5)   # w:sz=19
FONT_SIZE_BODY     = Pt(9)     # w:sz=18
FONT_SIZE_JOBTITLE = Pt(10)    # w:sz=20
TAB_STOP_RIGHT     = 10080     # twips — right edge of 7-inch text area (0.75-in margins)

OUT_DIR = Path("output")
OUT_DIR.mkdir(exist_ok=True)


# ══════════════════════════════════════════════════════════════════════════════
# Low-level helpers
# ══════════════════════════════════════════════════════════════════════════════

def _set_run_font(run, size: Pt, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = size
    # Only emit the tag when True — avoids <w:b w:val="0"/> / <w:i w:val="0"/> noise
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    # No color tag — ATS parsers prefer absence over explicit black


def _new_doc() -> Document:
    doc = Document()
    # Remove all default styles from Normal paragraph
    style = doc.styles["Normal"]
    style.font.name  = FONT
    style.font.size  = FONT_SIZE_BODY
    # Page setup: US Letter, 0.6-in top, 0.5-in bottom, 0.75-in left/right
    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.top_margin    = Twips(864)    # 0.6 inch
    sec.bottom_margin = Twips(720)    # 0.5 inch
    sec.left_margin   = Twips(1080)   # 0.75 inch
    sec.right_margin  = Twips(1080)   # 0.75 inch
    # Remove default paragraph spacing
    from docx.shared import Pt as _Pt
    style.paragraph_format.space_before = _Pt(0)
    style.paragraph_format.space_after  = _Pt(0)
    return doc


def _clear_default_spacing(doc: Document):
    """Kill Word's default 8pt-after-paragraph on every paragraph."""
    for p in doc.paragraphs:
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def _add_bottom_border(paragraph):
    """Add a thin bottom border line under a paragraph (section heading rule)."""
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")        # 0.75pt line
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "000000")
    pBdr.append(bot)
    pPr.append(pBdr)


def _add_right_tab(paragraph, pos: int = TAB_STOP_RIGHT):
    """Add a right-aligned tab stop at `pos` twips."""
    pPr  = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos))
    tabs.append(tab)
    pPr.append(tabs)


def _para(doc: Document, space_before=0, space_after=0):
    """Add a new paragraph with explicit spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p


# ══════════════════════════════════════════════════════════════════════════════
# Resume builder
# ══════════════════════════════════════════════════════════════════════════════

def _resume_section_heading(doc: Document, text: str):
    p = _para(doc, space_before=4, space_after=2)   # 80 DXA before
    run = p.add_run(text.upper())
    _set_run_font(run, FONT_SIZE_HEADING, bold=True)
    _add_bottom_border(p)
    return p


def _resume_name(doc: Document, name: str):
    p = _para(doc, space_before=0, space_after=2)
    p.alignment = 1  # CENTER
    run = p.add_run(name)
    _set_run_font(run, FONT_SIZE_NAME, bold=True)


def _resume_contact(doc: Document, contact: dict):
    parts = []
    for key in ("phone", "email", "linkedin", "github", "vercel", "location"):
        val = contact.get(key, "").strip()
        if val:
            parts.append(val)
    parts += [v for v in contact.get("other", []) if v.strip()]

    p = _para(doc, space_before=0, space_after=4)
    p.alignment = 1  # CENTER
    run = p.add_run("  |  ".join(parts))
    _set_run_font(run, FONT_SIZE_CONTACT)


def _resume_summary(doc: Document, summary: str):
    _resume_section_heading(doc, "SUMMARY")
    p = _para(doc, space_before=2, space_after=0)
    run = p.add_run(summary.strip())
    _set_run_font(run, FONT_SIZE_BODY)


def _resume_skills(doc: Document, skills: dict):
    _resume_section_heading(doc, "SKILLS")
    for category, items in skills.items():
        if not items:
            continue
        p = _para(doc, space_before=2, space_after=0)
        # Bold label
        label_run = p.add_run(f"{category}: ")
        _set_run_font(label_run, FONT_SIZE_BODY, bold=True)
        # Plain values
        val_run = p.add_run(", ".join(items))
        _set_run_font(val_run, FONT_SIZE_BODY)


def _resume_experience(doc: Document, experience: list):
    _resume_section_heading(doc, "EXPERIENCE")
    for entry in experience:
        # Header line: Bold title | Italic org \t Italic dates
        p = _para(doc, space_before=3, space_after=0)   # 60 DXA before
        _add_right_tab(p)

        title_run = p.add_run(entry["title"])
        _set_run_font(title_run, FONT_SIZE_JOBTITLE, bold=True)

        # Compose org string: company + optional location
        org_parts = [entry.get("company", "").strip()]
        loc = entry.get("location", "").strip()
        if loc:
            org_parts.append(loc)
        org_str = "  |  ".join(org_parts) if len(org_parts) > 1 else org_parts[0]

        sep_run = p.add_run("  |  ")
        _set_run_font(sep_run, FONT_SIZE_BODY, italic=True)

        org_run = p.add_run(org_str)
        _set_run_font(org_run, FONT_SIZE_BODY, italic=True)

        tab_run = p.add_run("\t")
        _set_run_font(tab_run, FONT_SIZE_BODY)

        date_run = p.add_run(entry.get("dates", ""))
        _set_run_font(date_run, FONT_SIZE_BODY, italic=True)

        # Bullets
        for bullet in entry.get("bullets", []):
            bp = _para(doc, space_before=0.8, space_after=0.8)  # 16 DXA each
            bp.paragraph_format.left_indent  = Inches(0.2)
            bp.paragraph_format.first_line_indent = Inches(-0.2)
            bullet_run = bp.add_run(f"• {bullet}")
            _set_run_font(bullet_run, FONT_SIZE_BODY)


def _resume_education(doc: Document, education: list):
    _resume_section_heading(doc, "EDUCATION")
    for entry in education:
        p = _para(doc, space_before=3, space_after=0)   # 60 DXA before
        _add_right_tab(p)

        degree_run = p.add_run(entry["degree"])
        _set_run_font(degree_run, FONT_SIZE_BODY, bold=True)

        school_str = entry.get("school", "").strip()
        gpa_str    = entry.get("gpa", "").strip()
        if school_str:
            sep_run = p.add_run("   -   ")
            _set_run_font(sep_run, FONT_SIZE_BODY, bold=True)
            school_run = p.add_run(school_str)
            _set_run_font(school_run, FONT_SIZE_BODY, italic=True)
        if gpa_str:
            gpa_run = p.add_run(f"  GPA: {gpa_str}")
            _set_run_font(gpa_run, FONT_SIZE_BODY, italic=True)

        tab_run = p.add_run("\t")
        _set_run_font(tab_run, FONT_SIZE_BODY)

        date_run = p.add_run(entry.get("dates", ""))
        _set_run_font(date_run, FONT_SIZE_BODY, italic=True)

        for bullet in entry.get("bullets", []):
            bp = _para(doc, space_before=0.8, space_after=0.8)  # 16 DXA each
            bp.paragraph_format.left_indent       = Inches(0.2)
            bp.paragraph_format.first_line_indent = Inches(-0.2)
            bullet_run = bp.add_run(f"- {bullet}")
            _set_run_font(bullet_run, FONT_SIZE_BODY)


def build_resume(resume_json_path: str | Path | dict, company: str) -> Path:
    """
    Build an ATS-clean resume .docx from a tailored resume JSON.
    Accepts either a file path (str/Path) or an already-loaded dict.
    Returns the output file path.
    """
    data    = resume_json_path if isinstance(resume_json_path, dict) \
              else json.loads(Path(resume_json_path).read_text())
    company = company.strip()
    outfile = OUT_DIR / f"AzatSh-{company}-resume.docx"

    doc = _new_doc()

    _resume_name(doc, data["name"])
    _resume_contact(doc, data["contact"])

    if data.get("summary"):
        _resume_summary(doc, data["summary"])

    if data.get("skills"):
        _resume_skills(doc, data["skills"])

    if data.get("experience"):
        _resume_experience(doc, data["experience"])

    if data.get("education"):
        _resume_education(doc, data["education"])

    doc.save(outfile)
    print(f"Resume  → {outfile}")
    return outfile


# ══════════════════════════════════════════════════════════════════════════════
# Cover letter builder
# ══════════════════════════════════════════════════════════════════════════════

def build_cover_letter(cover_text: str, company: str,
                       name: str = "Azat Shakirov",
                       email: str = "",
                       linkedin: str = "",
                       job_location: str = "") -> Path:
    """
    Build a cover letter .docx from plain cover letter text.
    Returns the output file path.
    """
    company = company.strip()
    outfile = OUT_DIR / f"AzatSh-{company}-coverL.docx"

    doc = _new_doc()
    sec = doc.sections[0]
    sec.top_margin    = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin   = Inches(1)
    sec.right_margin  = Inches(1)

    CL_SIZE    = Pt(11)
    CL_SPACING = 1.15  # line spacing multiplier

    def cl_para(text: str = "", bold=False, italic=False,
                space_before=0, space_after=6, align=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if align is not None:
            p.alignment = align
        # 1.15 line spacing
        p.paragraph_format.line_spacing = CL_SPACING
        if text:
            run = p.add_run(text)
            _set_run_font(run, CL_SIZE, bold=bold, italic=italic)
        return p

    # ── Header block ──────────────────────────────────────────────────────────
    cl_para(name, bold=True, space_after=2)
    contact_parts = [x for x in [email, linkedin] if x.strip()]
    if contact_parts:
        cl_para("  |  ".join(contact_parts), space_after=12)

    # ── Date ──────────────────────────────────────────────────────────────────
    cl_para(date.today().strftime("%B %d, %Y"), space_after=12)

    # ── Body paragraphs ───────────────────────────────────────────────────────
    # Defensive pre-processing: if Claude included a "---"-delimited header
    # section (contact + recipient block), discard everything up to and
    # including the last "---" marker, keeping only the body text.
    raw = cover_text.strip()
    _sep_re = re.compile(r'(?m)^\s*-{3,}\s*$')
    _sep_positions = [m.start() for m in _sep_re.finditer(raw)]
    if _sep_positions:
        # Drop everything through the final "---"
        last_sep = _sep_positions[-1]
        raw = raw[last_sep:].lstrip("-").strip()

    # Split on blank lines (double newline) first; fall back to single newline
    if "\n\n" in raw:
        paragraphs = [p.strip() for p in raw.split("\n\n") if p.strip()]
    else:
        paragraphs = [p.strip() for p in raw.split("\n") if p.strip()]

    # If no "---" were present, still strip a leading contact-header paragraph
    # (detected by: contains "|" and the sender name or email)
    if paragraphs and not _sep_positions:
        first = paragraphs[0]
        name_first = name.lower().split()[0]
        if "|" in first and (name_first in first.lower() or email.lower() in first.lower()):
            paragraphs.pop(0)

    # Strip trailing "sincerely/best/…" and/or trailing name paragraph —
    # we'll add the closing block ourselves to ensure initials are always present
    _closing_words = ("sincerely", "best", "regards", "thank you", "yours")
    if paragraphs and paragraphs[-1].strip().lower() == name.strip().lower():
        paragraphs.pop()
    if paragraphs and any(paragraphs[-1].lower().startswith(w) for w in _closing_words):
        paragraphs.pop()

    for para_text in paragraphs:
        cl_para(para_text, space_after=8)

    # ── Relocation line ───────────────────────────────────────────────────────
    if job_location and job_location.lower() not in ("remote", ""):
        reloc = (
            f"I'm currently based in Illinois completing my CS degree at Knox College "
            f"and am fully available to relocate to {job_location} for the duration of "
            f"the internship."
        )
        cl_para(reloc, space_after=8)

    # ── Closing block ─────────────────────────────────────────────────────────
    cl_para("Sincerely,", space_before=8, space_after=0)
    cl_para(name, space_after=0)

    doc.save(outfile)
    print(f"Cover letter → {outfile}")
    return outfile


# ══════════════════════════════════════════════════════════════════════════════
# CLI entry point
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python build_docs.py <resume_json> <company_name> [cover_letter_text]")
        sys.exit(1)

    resume_json = sys.argv[1]
    company     = sys.argv[2]
    cover_text  = sys.argv[3] if len(sys.argv) > 3 else None

    resume_data = json.loads(Path(resume_json).read_text())
    contact     = resume_data.get("contact", {})

    resume_out = build_resume(resume_json, company)

    if cover_text:
        cl_out = build_cover_letter(
            cover_text, company,
            name     = resume_data.get("name", ""),
            email    = contact.get("email", ""),
            linkedin = contact.get("linkedin", ""),
        )
